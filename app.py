import streamlit as st
import pdfplumber
from docx import Document
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import re
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from groq import Groq
from dotenv import load_dotenv

# ---------------- LOAD ENV ----------------
load_dotenv()

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="ATS Resume Optimizer",
    page_icon="🤖",
    layout="centered"
)

# ---------------- LOAD MODEL ----------------
@st.cache_resource
def load_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

model = load_model()

# ---------------- GROQ CLIENT ----------------
client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
)

# ---------------- CLEAN TEXT ----------------
def clean_text(text):

    text = re.sub(r"\(cid:\d+\)", " ", text)
    text = re.sub(r"\s+", " ", text)

    return text.strip()

# ---------------- EXTRACT TEXT ----------------
def extract_text(file):

    text = ""

    if file.name.endswith(".pdf"):

        with pdfplumber.open(file) as pdf:

            text = " ".join(
                [page.extract_text() or "" for page in pdf.pages]
            )

    elif file.name.endswith(".docx"):

        doc = Document(file)

        text = " ".join(
            [para.text for para in doc.paragraphs]
        )

    return clean_text(text)

# ---------------- KEYWORD EXTRACTION ----------------
def extract_keywords(text):

    words = re.findall(r"\b[a-zA-Z]{3,}\b", text.lower())

    stopwords = {
        "the","and","for","with","you","are","this","that","will",
        "have","from","your","using","based","into","within","their",
        "they","them","who","what","when","where","while","been",
        "being","also","than","then","such","both","through"
    }

    keywords = [
        word for word in words
        if word not in stopwords
    ]

    return list(set(keywords))

# ---------------- KEYWORD SCORE ----------------
def keyword_score(resume, job):

    job_keywords = set(extract_keywords(job))
    resume_keywords = set(extract_keywords(resume))

    matched = job_keywords & resume_keywords
    missing = job_keywords - resume_keywords

    score = (
        len(matched) / len(job_keywords)
        if job_keywords else 0
    )

    return score, matched, missing

# ---------------- SKILLS EXTRACTION ----------------
def extract_skills_section(text):

    text = text.lower()

    if "skills" in text:
        return text.split("skills")[1][:700]

    return text

# ---------------- SKILLS SCORE ----------------
def skills_score(resume, job):

    resume_skills = set(
        extract_keywords(
            extract_skills_section(resume)
        )
    )

    job_skills = set(
        extract_keywords(
            extract_skills_section(job)
        )
    )

    score = (
        len(resume_skills & job_skills) / len(job_skills)
        if job_skills else 0
    )

    return score

# ---------------- FINAL ATS SCORE ----------------
def get_score(resume, job):

    embeddings = model.encode([resume, job])

    semantic = cosine_similarity(
        [embeddings[0]],
        [embeddings[1]]
    )[0][0]

    keyword, matched, missing = keyword_score(
        resume,
        job
    )

    skills = skills_score(
        resume,
        job
    )

    final_score = (
        0.4 * keyword +
        0.4 * semantic +
        0.2 * skills
    )

    return {
        "final": final_score,
        "keyword": keyword,
        "semantic": semantic,
        "skills": skills,
        "matched": matched,
        "missing": missing
    }

# ---------------- GROQ AI CALL ----------------
def call_llm(prompt):

    try:

        # Prevent token overflow
        prompt = prompt[:6000]

        completion = client.chat.completions.create(

            model="llama3-8b-8192",

            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],

            temperature=0.4,

            max_tokens=1000
        )

        return completion.choices[0].message.content

    except Exception as e:

        return f"Error generating AI response: {str(e)}"

# ---------------- IMPROVE RESUME ----------------
def improve_resume(resume, job):

    prompt = f"""
You are a professional ATS Resume Optimizer.

TASK:
Improve the resume for the given job description.

STRICT RULES:
- Keep original structure
- Keep all technical skills
- Improve weak bullet points
- Add measurable impact
- Improve ATS keyword alignment
- Do NOT hallucinate fake experience
- Keep formatting professional

RETURN:
1. Improved Resume
2. Improvements Summary

JOB DESCRIPTION:
{job[:2500]}

RESUME:
{resume[:2500]}
"""

    return call_llm(prompt)

# ---------------- SAVE DOCX ----------------
def save_docx(text):

    doc = Document()

    for line in text.split("\n"):

        doc.add_paragraph(line)

    path = "improved_resume.docx"

    doc.save(path)

    return path

# ---------------- SAVE PDF ----------------
def save_pdf(text):

    path = "improved_resume.pdf"

    c = canvas.Canvas(
        path,
        pagesize=letter
    )

    y = 760

    for line in text.split("\n"):

        line = line[:100]

        c.drawString(40, y, line)

        y -= 15

        if y < 40:

            c.showPage()

            y = 760

    c.save()

    return path

# ---------------- UI ----------------
st.title("🤖 Recruiter-Level ATS Resume Analyzer")

st.markdown(
    """
Upload your resume and paste a job description to:

- Analyze ATS match score
- Detect missing keywords
- Improve resume using AI
- Download improved DOCX/PDF
"""
)

uploaded_file = st.file_uploader(
    "📄 Upload Resume",
    type=["pdf", "docx"]
)

job_description = st.text_area(
    "💼 Paste Job Description",
    height=220
)

# ---------------- MAIN LOGIC ----------------
if uploaded_file and job_description:

    with st.spinner("Analyzing resume..."):

        resume_text = extract_text(uploaded_file)

        result = get_score(
            resume_text,
            job_description
        )

    # ---------------- SCORE ----------------
    st.subheader("📊 ATS Match Score")

    st.metric(
        "Overall Match",
        f"{round(result['final'] * 100, 2)}%"
    )

    # ---------------- BREAKDOWN ----------------
    st.subheader("📈 Score Breakdown")

    st.write(
        f"✅ Keyword Match: {round(result['keyword'] * 100, 2)}%"
    )

    st.write(
        f"✅ Semantic Match: {round(result['semantic'] * 100, 2)}%"
    )

    st.write(
        f"✅ Skills Match: {round(result['skills'] * 100, 2)}%"
    )

    # ---------------- MATCHED ----------------
    st.subheader("✅ Matched Keywords")

    matched_keywords = list(result["matched"])[:30]

    st.write(matched_keywords)

    # ---------------- MISSING ----------------
    st.subheader("❌ Missing Keywords")

    missing_keywords = list(result["missing"])[:30]

    st.write(missing_keywords)

    # ---------------- IMPROVE BUTTON ----------------
    if st.button("✨ Improve Resume with AI"):

        with st.spinner("Generating improved resume..."):

            improved_resume = improve_resume(
                resume_text,
                job_description
            )

        st.subheader("🧠 Improved Resume")

        st.text_area(
            "",
            improved_resume,
            height=450
        )

        # ---------------- EXPORT FILES ----------------
        docx_path = save_docx(improved_resume)

        pdf_path = save_pdf(improved_resume)

        # ---------------- DOWNLOAD DOCX ----------------
        with open(docx_path, "rb") as file:

            st.download_button(
                label="⬇ Download DOCX",
                data=file,
                file_name="improved_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # ---------------- DOWNLOAD PDF ----------------
        with open(pdf_path, "rb") as file:

            st.download_button(
                label="⬇ Download PDF",
                data=file,
                file_name="improved_resume.pdf",
                mime="application/pdf"
            )
